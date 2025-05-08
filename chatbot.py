from fastapi import FastAPI, Request, UploadFile, File
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
import os
import requests
from bs4 import BeautifulSoup
import chromadb
from chromadb.utils import embedding_functions
import re
from typing import List, Dict, Optional
import logging
import PyPDF2
import docx
import io
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(title="Support Assistant",
             description="A RAG-based chatbot that provides information from customer support documentation")

# Serve static files
app.mount("/static", StaticFiles(directory="."), name="static")

# Pydantic models for request/response
class ChatRequest(BaseModel):
    message: str

class ChatResponse(BaseModel):
    response: str

class ScrapedContent(BaseModel):
    text: str
    source: str

class DocumentContent(BaseModel):
    text: str
    source: str
    
    filename: str

class DocumentProcessor:
    @staticmethod
    def process_pdf(file_content: bytes, filename: str) -> List[DocumentContent]:
        """Process PDF file and extract text content."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            contents = []
            
            logger.info(f"Processing PDF file: {filename} with {len(pdf_reader.pages)} pages")
            
            for page_num in range(len(pdf_reader.pages)):
                try:
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        contents.append(DocumentContent(
                            text=text.strip(),
                            source='PDF Document',
                            filename=f"{filename} (Page {page_num + 1})"
                        ))
                        logger.info(f"Successfully extracted text from page {page_num + 1}")
                    else:
                        logger.warning(f"No text content found on page {page_num + 1}")
                except Exception as e:
                    logger.error(f"Error processing page {page_num + 1}: {str(e)}")
            
            if not contents:
                logger.warning(f"No content could be extracted from PDF: {filename}")
            else:
                logger.info(f"Successfully processed PDF: {filename} with {len(contents)} content chunks")
            return contents
        except Exception as e:
            logger.error(f"Error processing PDF {filename}: {str(e)}")
            return []

    @staticmethod
    def process_docx(file_content: bytes, filename: str) -> List[DocumentContent]:
        """Process DOCX file and extract text content."""
        try:
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            contents = []
            
            logger.info(f"Processing DOCX file: {filename}")
            
            # Process paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    contents.append(DocumentContent(
                        text=para.text.strip(),
                        source='DOCX Document',
                        filename=f"{filename} (Paragraph {i + 1})"
                    ))
            
            # Process tables
            for i, table in enumerate(doc.tables):
                for j, row in enumerate(table.rows):
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        contents.append(DocumentContent(
                            text=row_text,
                            source='DOCX Document Table',
                            filename=f"{filename} (Table {i + 1}, Row {j + 1})"
                        ))
            
            if not contents:
                logger.warning(f"No content could be extracted from DOCX: {filename}")
            else:
                logger.info(f"Successfully processed DOCX: {filename} with {len(contents)} content chunks")
            return contents
        except Exception as e:
            logger.error(f"Error processing DOCX {filename}: {str(e)}")
            return []

    @staticmethod
    def process_file(file: UploadFile) -> List[DocumentContent]:
        """Process uploaded file based on its type."""
        try:
            content = file.file.read()
            filename = file.filename
            
            logger.info(f"Processing file: {filename}")
            
            if filename.lower().endswith('.pdf'):
                return DocumentProcessor.process_pdf(content, filename)
            elif filename.lower().endswith(('.docx', '.doc')):
                return DocumentProcessor.process_docx(content, filename)
            else:
                logger.warning(f"Unsupported file type: {filename}")
                return []
        except Exception as e:
            logger.error(f"Error reading file {file.filename}: {str(e)}")
            return []

class ChromaDBManager:
    def __init__(self):
        try:
            self.client = chromadb.Client()
            self.embedding_function = embedding_functions.SentenceTransformerEmbeddingFunction(
                model_name="all-MiniLM-L6-v2"
            )
            self.collection = self._initialize_collection()
            logger.info("Successfully initialized ChromaDB client and collection")
        except Exception as e:
            logger.error(f"Error initializing ChromaDB: {str(e)}")
            raise

    def _initialize_collection(self):
        """Initialize or get existing ChromaDB collection."""
        try:
            # Try to delete existing collection
            try:
                self.client.delete_collection(name="support_docs")
                logger.info("Deleted existing ChromaDB collection")
            except:
                logger.info("No existing collection to delete")

            # Create new collection
            logger.info("Creating new ChromaDB collection")
            collection = self.client.create_collection(
                name="support_docs",
                embedding_function=self.embedding_function,
                metadata={"hnsw:space": "cosine"}  # Use cosine similarity for better matching
            )
            return collection
        except Exception as e:
            logger.error(f"Error initializing collection: {str(e)}")
            raise

    def add_documents(self, documents: List[str], metadatas: List[Dict], ids: List[str]):
        """Add documents to the collection."""
        try:
            if not documents:
                logger.warning("No documents to add")
                return

            logger.info(f"Adding {len(documents)} documents to ChromaDB")
            self.collection.add(
                documents=documents,
                metadatas=metadatas,
                ids=ids
            )
            logger.info("Successfully added documents to ChromaDB")
        except Exception as e:
            logger.error(f"Error adding documents to ChromaDB: {str(e)}")
            raise

    def query(self, query_text: str, n_results: int = 5) -> Dict:
        """Query the collection for relevant documents."""
        try:
            logger.info(f"Querying ChromaDB with: {query_text}")
            results = self.collection.query(
                query_texts=[query_text],
                n_results=n_results,
                include=["documents", "distances", "metadatas"]
            )
            
            if not results or not results['documents'] or not results['documents'][0]:
                logger.warning("No results found in ChromaDB")
                return None
                
            logger.info(f"Found {len(results['documents'][0])} relevant documents")
            return results
        except Exception as e:
            logger.error(f"Error querying ChromaDB: {str(e)}")
            raise

class ContentScraper:
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean text by removing extra whitespace and newlines."""
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    @staticmethod
    def scrape_angel_one_support() -> List[ScrapedContent]:
        """Scrape content from Angel One support page."""
        logger.info("Scraping Angel One support page...")
        url = "https://www.angelone.in/support"
        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')
        
        content = []
        
        # Extract Quick Links section
        quick_links = soup.find_all('div', class_='quick-links')
        for link in quick_links:
            title = link.find('h3')
            description = link.find('p')
            if title and description:
                content.append(ScrapedContent(
                    text=f"{title.text.strip()}: {description.text.strip()}",
                    source='Quick Links'
                ))
        
        # Extract other sections
        sections = soup.find_all(['h2', 'h3', 'p'])
        for section in sections:
            if section.text.strip():
                content.append(ScrapedContent(
                    text=ContentScraper.clean_text(section.text),
                    source='General Content'
                ))
        
        logger.info(f"Scraped {len(content)} content items")
        return content

class ChatbotService:
    def __init__(self):
        try:
            self.chroma_db = ChromaDBManager()
            self._initialize_knowledge_base()
            logger.info("Successfully initialized ChatbotService")
        except Exception as e:
            logger.error(f"Error initializing ChatbotService: {str(e)}")
            raise

    def _initialize_knowledge_base(self):
        """Initialize the knowledge base with scraped content."""
        try:
            logger.info("Loading content into ChromaDB...")
            content = ContentScraper.scrape_angel_one_support()
            
            if content:
                documents = [item.text for item in content]
                metadatas = [{'source': item.source} for item in content]
                ids = [f"doc_{i}" for i in range(len(documents))]
                
                self.chroma_db.add_documents(documents, metadatas, ids)
                logger.info(f"Successfully loaded {len(documents)} documents into ChromaDB")
            else:
                logger.warning("No content was scraped from the website")
        except Exception as e:
            logger.error(f"Error initializing knowledge base: {str(e)}")
            raise

    def add_documents(self, documents: List[DocumentContent]):
        """Add new documents to the knowledge base."""
        try:
            if documents:
                texts = [doc.text for doc in documents]
                metadatas = [{'source': doc.source, 'filename': doc.filename} for doc in documents]
                ids = [f"doc_{i}" for i in range(len(documents))]
                
                self.chroma_db.add_documents(texts, metadatas, ids)
                logger.info(f"Successfully added {len(documents)} new documents to ChromaDB")
            else:
                logger.warning("No documents to add")
        except Exception as e:
            logger.error(f"Error adding documents: {str(e)}")
            raise

    def get_response(self, query: str) -> str:
        """Get response for a user query."""
        logger.info(f"Received query: {query}")
        
        try:
            results = self.chroma_db.query(query.lower())
            
            if not results:
                logger.info("No relevant information found in knowledge base")
                return "I don't know. This information is not available in my knowledge base."
            
            documents = results['documents'][0]
            distances = results['distances'][0]
            metadatas = results['metadatas'][0]
            
            # Collect all relevant information
            relevant_info = []
            for doc, dist, meta in zip(documents, distances, metadatas):
                if dist < 0.8:  # Only include reasonably relevant content
                    source_info = f" (from {meta['source']}"
                    if 'filename' in meta:
                        source_info += f" - {meta['filename']}"
                    source_info += ")"
                    relevant_info.append(f"{doc}{source_info}")
            
            if relevant_info:
                # Combine relevant information into a comprehensive response
                if len(relevant_info) == 1:
                    response = relevant_info[0]
                else:
                    # Format multiple pieces of information
                    response = "Here's what I found:\n\n"
                    for i, info in enumerate(relevant_info, 1):
                        response += f"{i}. {info}\n"
                logger.info(f"Found {len(relevant_info)} relevant pieces of information")
            else:
                logger.info("No relevant information found within similarity threshold")
                response = "I don't know. This information is not available in my knowledge base."
                
        except Exception as e:
            logger.error(f"Error processing query: {str(e)}")
            response = "I'm sorry, I encountered an error while processing your request."
        
        logger.info(f"Response: {response}")
        return response

# Initialize chatbot service
chatbot = ChatbotService()

@app.get("/")
def root():
    """Serve the main page."""
    return FileResponse("index.html")

@app.post("/chat", response_model=ChatResponse)
def chat_endpoint(request: ChatRequest):
    """Handle chat requests."""
    return ChatResponse(response=chatbot.get_response(request.message))

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """Handle file uploads."""
    try:
        logger.info(f"Received file upload request for: {file.filename}")
        contents = DocumentProcessor.process_file(file)
        
        if contents:
            chatbot.add_documents(contents)
            logger.info(f"Successfully added {len(contents)} content chunks to knowledge base")
            return {"message": f"Successfully processed {file.filename} and added {len(contents)} content chunks to knowledge base"}
        else:
            logger.warning(f"No content could be extracted from {file.filename}")
            return {"message": f"No content could be extracted from {file.filename}"}
    except Exception as e:
        logger.error(f"Error processing file {file.filename}: {str(e)}")
        return {"message": f"Error processing {file.filename}: {str(e)}"}

# To run: uvicorn chatbot:app --host 0.0.0.0 --port 8199 